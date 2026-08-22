"""Generates the final report and PoC scripts.

Validation rule: the LLM writes prose, but every finding section is rendered
from the `findings` table directly (title/severity/status/evidence/poc), not
from LLM free text. The LLM's job is limited to: (a) an executive summary,
and (b) narrative descriptions for verified findings — and even those are
scanned for evidence_id citations; any sentence in the narrative that reads
as a new factual claim without a citation gets flagged, not trusted.
"""
from __future__ import annotations

import re
import stat
from pathlib import Path

from .evidence import EvidenceStore
from .llm import LLMConnector

CITATION_RE = re.compile(r"\[evidence:(\d+)\]")


def _narrative_for_finding(llm: LLMConnector, store: EvidenceStore, finding: dict) -> str:
    ev_records = [store.get_evidence(i) for i in finding["evidence_ids"]]
    ev_text = "\n\n".join(
        f"[evidence:{e.id}] tool={e.tool} command={e.command}\n{e.raw_output[:1500]}"
        for e in ev_records if e
    )
    prompt = f"""Write a short (3-6 sentence) technical description of this verified finding
for a pentest report. You MUST cite [evidence:ID] for every factual claim, using only the
evidence ids listed below. Do not state anything (including any flag, credential, or exact
output value) that isn't directly present in the cited evidence.

Finding: {finding['title']} (severity: {finding['severity']})
Evidence:
{ev_text}
"""
    try:
        text = llm.complete([{"role": "user", "content": prompt}])
    except Exception as e:  # noqa: BLE001
        text = f"(narrative generation failed: {e})"
    return text


def _strip_uncited_sentences(text: str, valid_ids: set[int], store: EvidenceStore, run_id: str) -> str:
    """Best-effort guard: sentences containing no [evidence:ID] citation, or
    citing an id not in this finding's evidence set, are removed and logged
    rather than shipped in the report as unverified fact."""
    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    kept = []
    for s in sentences:
        ids = {int(m) for m in CITATION_RE.findall(s)}
        if not ids:
            store.log_unverified_claim(run_id, s, "no evidence citation in report narrative")
            continue
        if not ids.issubset(valid_ids):
            store.log_unverified_claim(run_id, s, f"cited evidence {ids} not in finding's evidence set {valid_ids}")
            continue
        kept.append(s)
    return " ".join(kept) if kept else "(no citable narrative could be generated for this finding)"


def _poc_script(finding: dict) -> str:
    header = (
        "#!/usr/bin/env bash\n"
        "# SPY-3697 generated PoC — reproduction of a VERIFIED finding.\n"
        f"# Finding: {finding['title']} (severity: {finding['severity']})\n"
        f"# Grounded in evidence ids: {finding['evidence_ids']}\n"
        "# Only run this against systems you are authorized to test.\n"
        "set -euo pipefail\n\n"
    )
    body = finding.get("poc_command") or "echo 'No PoC command captured for this finding.'"
    return header + body + "\n"


def generate_report(
    store: EvidenceStore, run_id: str, target: str, workspace: Path,
    formats: list[str], llm: LLMConnector,
) -> list[Path]:
    findings = store.list_findings(run_id)
    verified = [f for f in findings if f["status"] == "verified"]
    candidates = [f for f in findings if f["status"] == "candidate"]
    rejected = [f for f in findings if f["status"] == "rejected"]

    poc_dir = workspace / "poc"
    poc_dir.mkdir(exist_ok=True)
    poc_paths = []
    sections = []

    for f in verified:
        valid_ids = set(f["evidence_ids"])
        narrative = _narrative_for_finding(llm, store, f)
        narrative = _strip_uncited_sentences(narrative, valid_ids, store, run_id)

        script_path = poc_dir / f"poc_finding_{f['id']}.sh"
        script_path.write_text(_poc_script(f))
        script_path.chmod(script_path.stat().st_mode | stat.S_IEXEC)
        poc_paths.append(script_path)

        sections.append(
            f"### [{f['severity'].upper()}] {f['title']} (finding #{f['id']}) — VERIFIED\n\n"
            f"{narrative}\n\n"
            f"**Evidence:** {', '.join(f'#{i}' for i in f['evidence_ids'])}\n\n"
            f"**PoC script:** `{script_path.relative_to(workspace)}`\n"
            f"```bash\n{f.get('poc_command','(none captured)')}\n```\n"
        )

    for f in candidates:
        sections.append(
            f"### [{f['severity'].upper()}] {f['title']} (finding #{f['id']}) — UNVERIFIED CANDIDATE\n\n"
            f"{f.get('description','')}\n\n"
            f"Not independently confirmed by re-verification; treat as a lead for manual follow-up.\n"
            f"**Supporting evidence:** {', '.join(f'#{i}' for i in f['evidence_ids'])}\n"
        )

    unverified_claims = store.conn.execute(
        "SELECT claim_text, reason FROM unverified_claims WHERE run_id=?", (run_id,)
    ).fetchall()

    md = [
        f"# SPY-3697 Report — {target}",
        f"Run ID: `{run_id}`",
        "",
        "## Summary",
        f"- Verified findings: {len(verified)}",
        f"- Unverified candidates: {len(candidates)}",
        f"- Rejected (failed re-verification): {len(rejected)}",
        f"- Claims discarded for lacking evidence: {len(unverified_claims)}",
        "",
        "## Findings",
        *sections,
    ]
    if unverified_claims:
        md.append("## Appendix: Discarded / Unverified Claims (excluded from findings above)")
        for text, reason in unverified_claims:
            md.append(f"- ~~{text}~~ — _reason: {reason}_")

    report_text = "\n".join(md)
    out_paths = []

    if "md" in formats:
        md_path = workspace / "report.md"
        md_path.write_text(report_text)
        out_paths.append(md_path)

    if "docx" in formats:
        docx_path = workspace / "report.docx"
        _write_docx(docx_path, target, run_id, verified, candidates, unverified_claims)
        out_paths.append(docx_path)

    return out_paths + poc_paths


def _write_docx(path: Path, target: str, run_id: str, verified: list[dict],
                 candidates: list[dict], unverified_claims: list) -> None:
    from docx import Document  # local import, optional dependency

    doc = Document()
    doc.add_heading(f"SPY-3697 Report — {target}", level=0)
    doc.add_paragraph(f"Run ID: {run_id}")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(f"Verified findings: {len(verified)}")
    doc.add_paragraph(f"Unverified candidates: {len(candidates)}")
    doc.add_paragraph(f"Claims discarded for lacking evidence: {len(unverified_claims)}")

    doc.add_heading("Verified Findings", level=1)
    for f in verified:
        doc.add_heading(f"[{f['severity'].upper()}] {f['title']} (#{f['id']})", level=2)
        doc.add_paragraph(f.get("description", ""))
        doc.add_paragraph(f"Evidence: {', '.join(f'#{i}' for i in f['evidence_ids'])}")
        doc.add_paragraph(f"PoC command: {f.get('poc_command','(none)')}")

    doc.add_heading("Unverified Candidates", level=1)
    for f in candidates:
        doc.add_heading(f"[{f['severity'].upper()}] {f['title']} (#{f['id']})", level=2)
        doc.add_paragraph(f.get("description", ""))

    doc.save(path)
